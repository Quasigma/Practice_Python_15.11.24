# from openpyxl import load_workbook
# from openpyxl import Workbook

# wb = load_workbook('practice.xlsx')
# sheet = wb.active
# row1 = [cell.value for cell in sheet[1]]
# row2 = [cell.value for cell in sheet[2]]
# print(row1, row2)

# new_workbook = Workbook()
# new_sheet = new_workbook.active

# for i, value in enumerate(row1, start=1):
#     new_sheet.cell(row=1, column=1, value=value)

# for i, value in enumerate(row2,start=1):
#     new_sheet.cell(row=i, column=2,value=value)

# new_workbook.save("vertical_practice.xlsx")

#Json
# import json
# import requests

# url = 'https://jsonplaceholder.typicode.com/todos/1'
# response = requests.get(url)
# data = response.json()

# with open("practice_file.json", "w") as json_file:
#     json.dump(data, json_file)

# # Read Json

# with open("practice_file.json", "r") as json_file:
#     json_read = json.load(json_file)
# print(json_read)

# Word
import docx

doc = docx.Document("practice_file.docx")
print(doc.paragraphs[0].text)

doc.add_paragraph("Hello Discord")
doc.save("practice_file2.docx")